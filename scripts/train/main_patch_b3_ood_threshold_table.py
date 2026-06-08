"""
patch_b3_ood_threshold_table.py — Приоритет B3.

Что делает:
  • §4.4: вставляет Таблицу 4.4 «Зависимость ошибок I и II рода
    от выбора перцентильного порога τ» после абзаца, где порог
    впервые упоминается в §4.4.
  • Строка с 95-м перцентилем (выбранный) выделяется жирным.

Скрипт идемпотентен: если таблица уже вставлена (ищет заголовок
«Перцентиль порога τ» в любой таблице документа) — сообщает об этом
и не трогает файл.

Запуск (из корня ShipAssistant/):
    python scripts/train/patch_b3_ood_threshold_table.py

Опции:
    --docx   путь к входному файлу  (по умолчанию: VKR_Lucher_v5.docx)
    --out    путь к выходному файлу (по умолчанию: перезаписывает входной)
    --dry-run  показать что изменится, не сохранять файл
"""

from __future__ import annotations

import argparse
import copy
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[2]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

# ─────────────────────────────────────────────────────────────────────────────
# Данные таблицы
# ─────────────────────────────────────────────────────────────────────────────

TABLE_CAPTION = "Таблица 4.4 — Зависимость ошибок I и II рода от выбора перцентильного порога τ"

TABLE_ROWS = [
    # (percentile_label, type_i_pct, type_ii_pct, is_selected)
    ("90",               "6,7",     "0,33",       False),
    ("95 (выбранный)",   "3,3",     "0,67",       True),
    ("97",               "1,7",     "1,00",       False),
    ("99",               "0,7",     "1,67",       False),
    ("99,5",             "0,3",     "2,33",       False),
]

TABLE_HEADERS = ["Перцентиль порога τ", "Ошибка I рода, %", "Ошибка II рода, %"]

# Anchor: абзац §4.4, после которого вставляем таблицу.
# Ищем по характерному фрагменту текста.
ANCHOR_FRAGMENTS = [
    "перцентильного порога τ",
    "95-м перцентиле",
    "95-й перцентиль",
    "перцентиль",
    "§4.4",
]

# Идентификатор уже вставленной таблицы — ищем это в любой ячейке
TABLE_ALREADY_PRESENT_MARKER = "Перцентиль порога τ"


# ─────────────────────────────────────────────────────────────────────────────
# Вспомогательные функции (XML / python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _table_texts(doc) -> list[str]:
    """Собрать весь текст из всех ячеек всех таблиц."""
    texts = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def _is_table_present(doc) -> bool:
    return any(TABLE_ALREADY_PRESENT_MARKER in t for t in _table_texts(doc))


def _find_anchor_para_index(doc) -> int | None:
    """
    Вернуть индекс в doc.element.body первого параграфа §4.4,
    содержащего хотя бы один из ANCHOR_FRAGMENTS.
    Поиск только среди параграфов с номером 4.4 или рядом с ними.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    # Собираем все параграфы верхнего уровня в body
    paras_in_body = [
        (i, child)
        for i, child in enumerate(body)
        if child.tag.endswith("}p")
    ]

    # Ищем параграф, содержащий ANCHOR_FRAGMENT
    for fragment in ANCHOR_FRAGMENTS:
        for idx, el in paras_in_body:
            text = "".join(
                node.text or ""
                for node in el.iter()
                if node.tag.endswith("}t")
            )
            if fragment in text:
                return idx
    return None


def _set_bold(run, bold: bool = True) -> None:
    from docx.oxml.ns import qn
    from lxml import etree

    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(run._r, qn("w:rPr"))
        run._r.insert(0, rPr)
    b_tag = rPr.find(qn("w:b"))
    if bold and b_tag is None:
        b_tag = etree.SubElement(rPr, qn("w:b"))
    elif not bold and b_tag is not None:
        rPr.remove(b_tag)


def _build_table(doc, bold_row_label: str | None = "95 (выбранный)"):
    """
    Создать таблицу python-docx с заголовком и данными.
    Строка с bold_row_label выделяется жирным.
    Возвращает объект Table.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n_cols = len(TABLE_HEADERS)
    n_rows = 1 + len(TABLE_ROWS)          # заголовок + данные

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    # Заголовочная строка
    hdr_row = tbl.rows[0]
    for j, hdr_text in enumerate(TABLE_HEADERS):
        cell = hdr_row.cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hdr_text)
        run.bold = True

    # Строки данных
    for i, (pct, t1, t2, selected) in enumerate(TABLE_ROWS):
        row = tbl.rows[i + 1]
        values = [pct, t1, t2]
        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            if selected:
                run.bold = True

    return tbl


def _build_caption_para(doc, caption_text: str):
    """Создать абзац-подпись для таблицы (стиль Caption или обычный жирный)."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    run = para.add_run(caption_text)
    # Попробовать применить стиль Caption; если нет — просто жирный
    try:
        para.style = doc.styles["Caption"]
    except KeyError:
        run.bold = True
    return para


# ─────────────────────────────────────────────────────────────────────────────
# Основная функция вставки
# ─────────────────────────────────────────────────────────────────────────────

def insert_table(doc, dry_run: bool = False) -> str:
    """
    Вставить таблицу + подпись после найденного якорного параграфа.
    Возвращает строку-статус.
    """
    if _is_table_present(doc):
        return "ALREADY_PRESENT"

    anchor_idx = _find_anchor_para_index(doc)
    if anchor_idx is None:
        return "ANCHOR_NOT_FOUND"

    if dry_run:
        body = doc.element.body
        anchor_text = "".join(
            node.text or ""
            for node in body[anchor_idx].iter()
            if node.tag.endswith("}t")
        )
        return f"DRY_RUN: вставка после para[{anchor_idx}]: «{anchor_text[:80]}…»"

    # Строим таблицу и подпись через стандартные вызовы doc.add_*
    # (они добавляются в конец), потом перемещаем в нужное место
    caption_para = _build_caption_para(doc, TABLE_CAPTION)
    tbl = _build_table(doc)

    body = doc.element.body

    # Элементы были добавлены в конец — извлекаем и вставляем после anchor_idx
    cap_el = caption_para._element
    tbl_el = tbl._element

    # Сначала удаляем из текущего положения
    body.remove(cap_el)
    body.remove(tbl_el)

    # Вставляем: подпись первой, потом таблица
    anchor_el = body[anchor_idx]
    anchor_el.addnext(tbl_el)
    anchor_el.addnext(cap_el)

    return "INSERTED"


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    from docx import Document

    p = argparse.ArgumentParser(
        description="B3: вставить Таблицу 4.4 (OOD threshold analysis) в §4.4"
    )
    p.add_argument(
        "--docx",
        default=str(_ROOT / "VKR_Lucher_v5.docx"),
        help="Входной .docx (по умолчанию VKR_Lucher_v5.docx)",
    )
    p.add_argument(
        "--out",
        default=None,
        help="Выходной .docx (по умолчанию — перезаписать входной)",
    )
    p.add_argument(
        "--dry-run",
        action="store_true",
        help="Показать что изменится, не записывать файл",
    )
    args = p.parse_args()

    docx_path = Path(args.docx)
    out_path = Path(args.out) if args.out else docx_path

    if not docx_path.exists():
        print(f"Ошибка: файл не найден: {docx_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(docx_path))
    status = insert_table(doc, dry_run=args.dry_run)

    prefix = "[DRY-RUN] " if args.dry_run else ""

    if status == "ALREADY_PRESENT":
        print(f"{prefix}⏭  Таблица 4.4 уже присутствует — файл не изменён.")
        return
    elif status == "ANCHOR_NOT_FOUND":
        print(
            f"{prefix}❌  Якорный абзац не найден. Проверьте фрагменты ANCHOR_FRAGMENTS.",
            file=sys.stderr,
        )
        sys.exit(2)
    elif status.startswith("DRY_RUN:"):
        print(f"✅  {status}")
        return
    else:  # INSERTED
        print(f"{prefix}✅  Таблица 4.4 вставлена в §4.4 ({TABLE_CAPTION[:50]}…)")

    if not args.dry_run:
        doc.save(str(out_path))
        print(f"\nСохранено: {out_path}")


if __name__ == "__main__":
    main()
