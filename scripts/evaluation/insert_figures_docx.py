"""
scripts/vkr/insert_figures_docx.py
====================================
Вставляет 8 рисунков в VKR_Lucher_v13.docx и сохраняет как VKR_Lucher_v14.docx.

Логика вставки каждого рисунка:
  1. Находим целевой параграф по индексу и тексту (двойная проверка).
  2. Вставляем параграф-ссылку (если ещё нет) + параграф с изображением + подпись.
  3. Обновляем строку реферата (количество рисунков).

Запуск:
    python scripts/vkr/insert_figures_docx.py
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------
PROJ = Path(__file__).resolve().parent.parent.parent
FIGS = PROJ / "artifacts" / "plots" / "vkr_figures"
SRC_DOCX = PROJ / "VKR_Lucher_v13.docx"
DST_DOCX = PROJ / "VKR_Lucher_v14.docx"


# ---------------------------------------------------------------------------
# Вспомогательные функции
# ---------------------------------------------------------------------------

def insert_paragraph_after(ref_para, doc: Document):
    """Создаёт новый параграф сразу после ref_para и возвращает его."""
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._element.getparent())


def insert_image_paragraph(ref_para, doc: Document, image_path: Path,
                            width_cm: float = 14.5):
    """Вставляет параграф с изображением после ref_para. Возвращает параграф."""
    img_para = insert_paragraph_after(ref_para, doc)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return img_para


def insert_caption(ref_para, doc: Document, caption_text: str):
    """Вставляет подпись под рисунком после ref_para. Возвращает параграф."""
    cap_para = insert_paragraph_after(ref_para, doc)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption_text)
    run.font.size = Pt(12)
    # Отступ сверху/снизу (6 пт)
    ppr = cap_para._element.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), "60")
    spc.set(qn("w:after"), "120")
    ppr.append(spc)
    return cap_para


def insert_ref_sentence(ref_para, doc: Document, sentence: str):
    """Вставляет параграф-ссылку (курсив, 10 пт) после ref_para."""
    s_para = insert_paragraph_after(ref_para, doc)
    run = s_para.add_run(sentence)
    run.font.size = Pt(10)
    run.font.italic = True
    return s_para


def insert_figure_block(after_para, doc: Document,
                         image_path: Path, caption: str,
                         ref_sentence: str | None = None,
                         width_cm: float = 14.5) -> tuple:
    """
    Вставляет:
      1. (опционально) параграф-ссылку
      2. Параграф с изображением
      3. Параграф с подписью

    Возвращает (ref_para | None, img_para, cap_para).
    """
    # Порядок insertAfter: каждый следующий addnext сдвигает предыдущий вниз
    # Чтобы получить порядок ref_sentence → image → caption,
    # вставляем в обратном порядке: сначала caption, затем image, затем ref.

    cap_para = insert_caption(after_para, doc, caption)
    img_para = insert_image_paragraph(after_para, doc, image_path, width_cm)

    if ref_sentence:
        ref_para = insert_ref_sentence(after_para, doc, ref_sentence)
    else:
        ref_para = None

    return ref_para, img_para, cap_para


def find_para(doc: Document, idx: int, expected_substr: str | None = None):
    """
    Возвращает параграф по индексу с опциональной проверкой содержимого.
    Поднимает ValueError, если текст не совпадает.
    """
    paras = doc.paragraphs
    para = paras[idx]
    if expected_substr and expected_substr not in para.text:
        raise ValueError(
            f"Параграф P{idx} не содержит ожидаемую строку '{expected_substr}'.\n"
            f"Фактический текст: '{para.text[:120]}'"
        )
    return para


# ---------------------------------------------------------------------------
# Вставка рисунков
# ---------------------------------------------------------------------------

def insert_all_figures(doc: Document) -> list[dict]:
    """
    Вставляет 8 рисунков по нужным позициям.
    Возвращает список с информацией о каждом рисунке.
    """
    inserted = []

    # ------------------------------------------------------------------
    # Рисунок 2.1 — Общая архитектура системы (§2.2, после P200)
    # ------------------------------------------------------------------
    p200 = find_para(doc, 200, "Разделение вычислительного процесса")
    ref2_1, img2_1, cap2_1 = insert_figure_block(
        p200, doc,
        FIGS / "fig_2_1_architecture.png",
        "Рисунок 2.1 — Общая архитектура системы распознавания голосовых команд",
        ref_sentence="Общая архитектура конвейера приведена на рисунке 2.1.",
        width_cm=14.5,
    )
    inserted.append({"num": "2.1", "title": "Архитектура системы", "stub": False})
    print("  ✓ Рисунок 2.1 вставлен после P200")

    # ------------------------------------------------------------------
    # Рисунок 2.2 — Пайплайн HybridAudioEngine (§2.7, после P236)
    # ------------------------------------------------------------------
    p236 = find_para(doc, 236, "Stage 3: Классификация намерения")
    _, _, _ = insert_figure_block(
        p236, doc,
        FIGS / "fig_2_2_hybrid_pipeline.png",
        "Рисунок 2.2 — Пайплайн HybridAudioEngine (пять стадий обработки)",
        ref_sentence="Пайплайн HybridAudioEngine показан на рисунке 2.2.",
        width_cm=14.5,
    )
    inserted.append({"num": "2.2", "title": "HybridAudioEngine", "stub": False})
    print("  ✓ Рисунок 2.2 вставлен после P236")

    # ------------------------------------------------------------------
    # Рисунок 3.1 — Схема LoRA-адаптации (§3.3, после P263)
    # ------------------------------------------------------------------
    p263 = find_para(doc, 263, "позволяет оптимизатору обновлять")
    _, _, _ = insert_figure_block(
        p263, doc,
        FIGS / "fig_3_1_lora_schema.png",
        "Рисунок 3.1 — Схема LoRA-адаптации слоя self-attention (r = 32, α = 64)",
        ref_sentence="Схема LoRA-адаптации приведена на рисунке 3.1.",
        width_cm=10.0,
    )
    inserted.append({"num": "3.1", "title": "Схема LoRA-адаптации", "stub": False})
    print("  ✓ Рисунок 3.1 вставлен после P263")

    # ------------------------------------------------------------------
    # Рисунок 4.2 — Кривые обучения (§4.2, после P299)
    # ------------------------------------------------------------------
    p299 = find_para(doc, 299, "Примечание: приведённые значения F1")
    _, _, _ = insert_figure_block(
        p299, doc,
        FIGS / "fig_4_2_training_curves.png",
        "Рисунок 4.2 — Кривые обучения LoRA-адаптации: Loss и macro-F1 по эпохам",
        ref_sentence=(
            "Динамика обучения в виде кривых потерь и метрики F1 "
            "по эпохам приведена на рисунке 4.2."
        ),
        width_cm=14.5,
    )
    inserted.append({"num": "4.2", "title": "Кривые обучения", "stub": False})
    print("  ✓ Рисунок 4.2 вставлен после P299")

    # ------------------------------------------------------------------
    # Рисунок 4.3 — Распределение Махаланобиса (§4.4, после P323)
    # ------------------------------------------------------------------
    p323 = find_para(doc, 323, "Таблица 4.4 демонстрирует")
    _, _, _ = insert_figure_block(
        p323, doc,
        FIGS / "fig_4_3_mahalanobis.png",
        "Рисунок 4.3 — Распределение расстояний Махаланобиса: "
        "целевые команды vs. внеклассовые сигналы",
        ref_sentence=(
            "Распределение расстояний Махаланобиса для целевых и внеклассовых "
            "образцов показано на рисунке 4.3."
        ),
        width_cm=12.0,
    )
    inserted.append({"num": "4.3", "title": "Распределение Махаланобиса",
                     "stub": False})
    print("  ✓ Рисунок 4.3 вставлен после P323")

    # ------------------------------------------------------------------
    # Рисунок 4.4 — Потребление RAM за 24 часа (§4.5, после P326)
    # ------------------------------------------------------------------
    p326 = find_para(doc, 326,
                     "Результаты 24-часового нагрузочного теста")
    _, _, _ = insert_figure_block(
        p326, doc,
        FIGS / "fig_4_4_ram_24h.png",
        "Рисунок 4.4 — Потребление оперативной памяти (RSS) за 24 часа "
        "нагрузочного теста",
        ref_sentence=(
            "График потребления памяти за 24 часа приведён на рисунке 4.4."
        ),
        width_cm=14.5,
    )
    inserted.append({"num": "4.4", "title": "RAM 24 ч", "stub": False})
    print("  ✓ Рисунок 4.4 вставлен после P326")

    # ------------------------------------------------------------------
    # Рисунок 4.5 — F1 vs SNR (§4.7, после P153)
    # P152 уже содержит ссылку на «рисунке 4.5»; вставляем рисунок после P153
    # ------------------------------------------------------------------
    p153 = find_para(doc, 153, "Предварительный анализ")
    _, _, _ = insert_figure_block(
        p153, doc,
        FIGS / "fig_4_5_f1_vs_snr.png",
        "Рисунок 4.5 — Деградация macro-F1 при снижении ОСШ "
        "(реальные измерения: чистая речь и SNR = 12 дБ)",
        ref_sentence=None,   # уже упомянуто в P152
        width_cm=12.0,
    )
    inserted.append({"num": "4.5", "title": "F1 vs SNR", "stub": False})
    print("  ✓ Рисунок 4.5 вставлен после P153")

    # ------------------------------------------------------------------
    # Рисунок 4.1 — Матрица ошибок (§5.2, после P345)
    # Вставляем ПОСЛЕДНИМ, чтобы не сдвигать индексы остальных разделов
    # ------------------------------------------------------------------
    p345 = find_para(doc, 345, "Наиболее частая ошибка")
    _, _, _ = insert_figure_block(
        p345, doc,
        FIGS / "fig_4_1_confusion_matrix.png",
        "Рисунок 4.1 — Матрица ошибок классификации (тестовая выборка, N = 123)",
        ref_sentence=(
            "Матрица ошибок классификации приведена на рисунке 4.1."
        ),
        width_cm=11.0,
    )
    inserted.append({"num": "4.1", "title": "Матрица ошибок", "stub": False})
    print("  ✓ Рисунок 4.1 вставлен после P345")

    return inserted


# ---------------------------------------------------------------------------
# Обновление реферата
# ---------------------------------------------------------------------------

def update_abstract(doc: Document) -> None:
    """
    Обновляет строку «содержит: N с., ...» в реферате, добавляя «8 рис.».
    """
    for para in doc.paragraphs[:40]:
        if "содержит:" in para.text and ("с." in para.text or "табл." in para.text):
            old_text = para.text
            # Добавляем «8 рис.» перед числом таблиц
            new_text = old_text.replace(
                "с., ", "с., 8 рис., ", 1
            )
            if new_text == old_text:
                # Запасной вариант: вставить в конец до точки
                new_text = old_text.rstrip(".") + ", 8 рис."
            # Заменяем текст в первом прогоне (сохраняем форматирование)
            if para.runs:
                full = "".join(r.text for r in para.runs)
                # Заменяем весь текст в первом прогоне
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.text = new_text
            print(f"  ✓ Реферат обновлён:\n"
                  f"    БЫЛО: {old_text}\n"
                  f"    СТАЛО: {new_text}")
            return
    print("  ⚠ Строка реферата не найдена — обновите вручную")


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Открываем {SRC_DOCX.name} …")
    doc = Document(str(SRC_DOCX))

    print("\nВставка рисунков:")
    inserted = insert_all_figures(doc)

    print("\nОбновление реферата:")
    update_abstract(doc)

    print(f"\nСохраняем → {DST_DOCX.name} …")
    doc.save(str(DST_DOCX))
    size_mb = DST_DOCX.stat().st_size / 1_048_576
    print(f"  ✓ Файл создан: {size_mb:.1f} МБ")

    print("\n" + "=" * 60)
    print("ИТОГО рисунков:")
    for item in inserted:
        tag = "ДАННЫЕ" if not item["stub"] else "ЗАГЛУШКА"
        print(f"  [{tag}]  Рисунок {item['num']} — {item['title']}")


if __name__ == "__main__":
    main()
